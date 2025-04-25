
"""
NFR Capture & Testing Strategy Assistant
Track Non-Functional Requirements with exportable documentation.
"""
import os
import yaml
import json
import tempfile
from docx import Document
from fpdf import FPDF
import markdown
import numpy as np
import faiss
from sentence_transformers import SentenceTransformer
from bs4 import BeautifulSoup
import requests

class NFRService:
    def __init__(self):
        # Load model for semantic search
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Sample NFR categories and their descriptions
        self.nfr_categories = {
            "performance": "Requirements related to response time, throughput, and resource utilization.",
            "security": "Requirements related to authentication, authorization, data protection, etc.",
            "reliability": "Requirements related to availability, fault tolerance, and recoverability.",
            "usability": "Requirements related to user experience, accessibility, and ease of use.",
            "maintainability": "Requirements related to code quality, documentation, and ease of maintenance.",
            "scalability": "Requirements related to handling growth in users, data, or transactions.",
            "compatibility": "Requirements related to interoperability with other systems or platforms.",
            "compliance": "Requirements related to regulatory, legal, or industry standards."
        }
        
        # Initialize index and embeddings for NFR categories
        self._create_embeddings()

    def _create_embeddings(self):
        """Create embeddings for NFR categories and descriptions"""
        texts = [f"{category}: {description}" for category, description in self.nfr_categories.items()]
        self.category_embeddings = self.embedding_model.encode(texts)
        
        # Create FAISS index for NFR categories
        embedding_dim = self.category_embeddings.shape[1]
        self.category_index = faiss.IndexFlatL2(embedding_dim)
        self.category_index.add(np.array(self.category_embeddings).astype('float32'))

    def suggest_nfr_category(self, requirement_text):
        """Suggest NFR category based on semantic search"""
        if not requirement_text:
            return None
            
        # Get query embedding
        query_embedding = self.embedding_model.encode([requirement_text])
        
        # Search in FAISS index
        distances, indices = self.category_index.search(np.array(query_embedding).astype('float32'), 3)
        
        suggestions = []
        categories = list(self.nfr_categories.keys())
        
        for i, idx in enumerate(indices[0]):
            if idx < len(categories):
                category = categories[idx]
                suggestions.append({
                    "category": category,
                    "description": self.nfr_categories[category],
                    "confidence": float(1 / (1 + distances[0][i]))
                })
        
        return suggestions

    def load_nfr_schema(self, schema_yaml):
        """Load NFR schema from YAML"""
        try:
            schema = yaml.safe_load(schema_yaml)
            return {"status": "success", "schema": schema}
        except Exception as e:
            return {"status": "error", "message": str(e)}

    def validate_nfr_data(self, nfr_data, schema):
        """Validate NFR data against schema"""
        # Basic validation - can be expanded
        for required_field in schema.get("required_fields", []):
            if required_field not in nfr_data or not nfr_data[required_field]:
                return {
                    "status": "error", 
                    "message": f"Missing required field: {required_field}"
                }
        
        return {"status": "success"}

    def generate_markdown_doc(self, nfr_data):
        """Generate markdown documentation from NFR data"""
        markdown_content = f"""# Non-Functional Requirements: {nfr_data.get('project_name', 'Untitled Project')}

## Project Overview
{nfr_data.get('project_description', 'No description provided')}

## NFR Categories
"""
        
        for category, requirements in nfr_data.get('requirements', {}).items():
            markdown_content += f"\n### {category.upper()}\n"
            
            for req in requirements:
                markdown_content += f"""
#### {req.get('id', 'UNKNOWN')} - {req.get('name', 'Untitled')}
- **Description**: {req.get('description', 'No description provided')}
- **Acceptance Criteria**: {req.get('acceptance_criteria', 'None specified')}
- **Priority**: {req.get('priority', 'Medium')}
- **Test Strategy**: {req.get('test_strategy', 'Not defined')}
"""
        
        if 'audit_trail' in nfr_data and nfr_data['audit_trail']:
            markdown_content += "\n## Audit Trail\n"
            for entry in nfr_data['audit_trail']:
                markdown_content += f"- {entry.get('timestamp', 'Unknown')}: {entry.get('action', 'Unknown action')} by {entry.get('user', 'Unknown user')}\n"
        
        return {"status": "success", "markdown": markdown_content}

    def export_to_docx(self, markdown_content):
        """Export markdown to DOCX format"""
        try:
            doc = Document()
            
            # Parse markdown content (simplified approach)
            lines = markdown_content.split('\n')
            
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], 0)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 1)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], 2)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], 3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='ListBullet')
                else:
                    if line.strip():
                        doc.add_paragraph(line)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
                doc.save(temp_file.name)
                temp_path = temp_file.name
            
            return {"status": "success", "file_path": temp_path}
            
        except Exception as e:
            return {"status": "error", "message": str(e)}

    def export_to_pdf(self, markdown_content):
        """Export markdown to PDF format"""
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Convert markdown to HTML for better parsing
            html_content = markdown.markdown(markdown_content)
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract text and format accordingly (simplified approach)
            for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'li']):
                if tag.name == 'h1':
                    pdf.set_font("Arial", 'B', 16)
                    pdf.cell(0, 10, tag.text, ln=True)
                elif tag.name == 'h2':
                    pdf.set_font("Arial", 'B', 14)
                    pdf.cell(0, 10, tag.text, ln=True)
                elif tag.name == 'h3':
                    pdf.set_font("Arial", 'B', 12)
                    pdf.cell(0, 10, tag.text, ln=True)
                elif tag.name == 'h4':
                    pdf.set_font("Arial", 'BI', 12)
                    pdf.cell(0, 10, tag.text, ln=True)
                elif tag.name == 'li':
                    pdf.set_font("Arial", size=12)
                    pdf.cell(0, 10, f"• {tag.text}", ln=True)
                else:  # p
                    pdf.set_font("Arial", size=12)
                    pdf.multi_cell(0, 10, tag.text)
                    
                pdf.ln(2)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                pdf.output(temp_file.name)
                temp_path = temp_file.name
            
            return {"status": "success", "file_path": temp_path}
            
        except Exception as e:
            return {"status": "error", "message": str(e)}

    def upload_to_confluence(self, page_id, markdown_content, base_url, auth_token):
        """Upload NFR documentation to Confluence"""
        headers = {
            "Authorization": f"Bearer {auth_token}",
            "Content-Type": "application/json"
        }
        
        # First, get the current version of the page
        version_url = f"{base_url}/rest/api/content/{page_id}"
        
        try:
            response = requests.get(version_url, headers=headers)
            response.raise_for_status()
            
            page_data = response.json()
            current_version = page_data.get("version", {}).get("number", 1)
            
            # Convert markdown to HTML
            html_content = markdown.markdown(markdown_content)
            
            # Update the page
            update_data = {
                "version": {
                    "number": current_version + 1
                },
                "title": page_data.get("title", "Non-Functional Requirements"),
                "type": "page",
                "body": {
                    "storage": {
                        "value": html_content,
                        "representation": "storage"
                    }
                }
            }
            
            update_url = f"{base_url}/rest/api/content/{page_id}"
            update_response = requests.put(
                update_url, 
                headers=headers, 
                data=json.dumps(update_data)
            )
            update_response.raise_for_status()
            
            return {"status": "success", "message": "NFR documentation uploaded to Confluence"}
            
        except Exception as e:
            return {"status": "error", "message": str(e)}
